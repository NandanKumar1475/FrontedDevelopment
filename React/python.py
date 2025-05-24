from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Define styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
doc.add_heading('Astha Tomar', 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Contact Info
doc.add_paragraph('📞 6203773120 | tomarastha2001@gmail.com | linkedin.com/in/AsthaTomar | github.com/AsthaTomar').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Experience
doc.add_heading('Experience', level=1)

doc.add_paragraph(
    'COGNIZANT TECHNOLOGY SERVICES, Chennai — Programmr Analyst Trainee\n'
    'Dec 2024\n'
    '- Undergoing Performance Engineering Training.\n'
    '- Skilled in LoadRunner and JMeter.\n'
    '- Proficient in AppDynamics, Perfmon, Core Java, AWS, Jenkins, Git, and more.\n'
)

doc.add_paragraph(
    'NOMURA SERVICES INDIA, Mumbai — Software Developer Intern\n'
    'Jan 2024 – June 2024\n'
    '- Developed servlets, JSPs, and used DAO pattern.\n'
    '- Improved data processing by 20%.\n'
    '- Used Spring, JUnit, Mockito, Jenkins, Git.\n'
)

doc.add_paragraph(
    'TATA STEEL, Jamshedpur — UI Developer Intern\n'
    'May 2023 – June 2023\n'
    '- Designed user-friendly interface; 15% increase in adoption.\n'
    '- Reduced response time by 1 hour.\n'
)

# Projects
doc.add_heading('Projects', level=1)

doc.add_paragraph(
    'Crowdfunding Decentralised Web Application\n'
    '- Built on Avalanche blockchain using Solidity.\n'
    '- Implemented smart contracts for secure transactions.\n'
)

doc.add_paragraph(
    'CRMS + UI Development\n'
    '- Static website using JavaScript, HTML/CSS, Bootstrap.\n'
    '- Improved user engagement by 30%.\n'
)

# Technical Skills
doc.add_heading('Technical Skills', level=1)
doc.add_paragraph(
    'Languages: Java 8, Python, SQL, JavaScript, JMeter, LoadRunner\n'
    'Frameworks: Spring Boot, MVC, JDBC\n'
    'Build Tools: Maven, Jenkins\n'
    'Monitoring: AppDynamics, Perfmon\n'
    'Version Control: Git, GitHub, GitLab\n'
    'Concepts: OOP, DSA, Networking, SDLC, AWS, API Design, CI/CD\n'
)

# Education
doc.add_heading('Education', level=1)
doc.add_paragraph(
    'Cochin University of Science and Technology, Kerala\n'
    'B.Tech in Computer Science (2020 - 2024) — GPA: 8.2\n'
)
doc.add_paragraph(
    'Rajendra Vidyalaya, Jamshedpur\n'
    'Matriculation: 94% (2018), 12th: 87% (2020)\n'
)

# Certifications
doc.add_heading('Certifications and Recognitions', level=1)
doc.add_paragraph(
    '• Introduction to Java (Coding Ninjas)\n'
    '• Tink-Her-Hack Hackathon — Top 3 among all teams (CODEMATES)\n'
)

# Save
output_path = "/mnt/data/AsthaTomar_Resume_Recreated.docx"
doc.save(output_path)

output_path
